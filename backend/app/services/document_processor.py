import os
import uuid
from typing import List, Dict, Any
from pathlib import Path
import logging

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
import aiofiles

from ..core.config import settings

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """文档处理服务，负责文档解析和文本分块"""
    
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
    
    async def process_file(self, file_path: str, filename: str) -> Dict[str, Any]:
        """处理上传的文件，返回文档信息和分块结果"""
        try:
            # 提取文本内容
            text_content = await self._extract_text(file_path, filename)
            
            if not text_content.strip():
                raise ValueError("文档内容为空")
            
            # 生成文档ID
            doc_id = str(uuid.uuid4())
            
            # 文本分块
            chunks = self._split_text(text_content, filename, doc_id)
            
            # 构建文档元数据
            document_info = {
                "id": doc_id,
                "title": filename,
                "content": text_content,
                "file_type": Path(filename).suffix.lower(),
                "source": filename,
                "chunk_count": len(chunks),
                "chunks": chunks
            }
            
            logger.info(f"成功处理文档: {filename}, 生成 {len(chunks)} 个文本块")
            return document_info
            
        except Exception as e:
            logger.error(f"处理文档失败 {filename}: {str(e)}")
            raise
    
    async def _extract_text(self, file_path: str, filename: str) -> str:
        """根据文件类型提取文本内容"""
        file_ext = Path(filename).suffix.lower()
        
        if file_ext == '.pdf':
            return await self._extract_pdf_text(file_path)
        elif file_ext == '.docx':
            return await self._extract_docx_text(file_path)
        elif file_ext in ['.txt', '.md']:
            return await self._extract_plain_text(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {file_ext}")
    
    async def _extract_pdf_text(self, file_path: str) -> str:
        """提取PDF文本"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                text_content = ""
                
                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n"
                
                return text_content.strip()
        except Exception as e:
            raise ValueError(f"PDF解析失败: {str(e)}")
    
    async def _extract_docx_text(self, file_path: str) -> str:
        """提取DOCX文本"""
        try:
            doc = DocxDocument(file_path)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
            
            return text_content.strip()
        except Exception as e:
            raise ValueError(f"DOCX解析失败: {str(e)}")
    
    async def _extract_plain_text(self, file_path: str) -> str:
        """提取纯文本"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                content = await file.read()
                return content.strip()
        except UnicodeDecodeError:
            # 尝试其他编码
            try:
                async with aiofiles.open(file_path, 'r', encoding='gbk') as file:
                    content = await file.read()
                    return content.strip()
            except Exception as e:
                raise ValueError(f"文本文件编码错误: {str(e)}")
        except Exception as e:
            raise ValueError(f"纯文本解析失败: {str(e)}")
    
    def _split_text(self, text: str, source: str, doc_id: str) -> List[Document]:
        """将文本分割成块"""
        # 使用LangChain的文本分割器
        text_chunks = self.text_splitter.split_text(text)
        
        documents = []
        for i, chunk in enumerate(text_chunks):
            doc = Document(
                page_content=chunk,
                metadata={
                    "source": source,
                    "document_id": doc_id,
                    "chunk_index": i,
                    "chunk_count": len(text_chunks),
                    "chunk_size": len(chunk)
                }
            )
            documents.append(doc)
        
        return documents
    
    def validate_file(self, filename: str, file_size: int) -> bool:
        """验证文件格式和大小"""
        # 检查文件扩展名
        file_ext = Path(filename).suffix.lower()
        if file_ext not in settings.allowed_extensions:
            raise ValueError(f"不支持的文件格式: {file_ext}")
        
        # 检查文件大小
        if file_size > settings.max_file_size:
            raise ValueError(f"文件大小超出限制: {file_size} bytes")
        
        return True
    
    async def cleanup_temp_file(self, file_path: str):
        """清理临时文件"""
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                logger.info(f"已清理临时文件: {file_path}")
        except Exception as e:
            logger.warning(f"清理临时文件失败 {file_path}: {str(e)}") 